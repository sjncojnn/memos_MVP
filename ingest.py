"""
Tiếp nhận & chuẩn hoá tri thức (Problem Statement mục 4):
  - Đọc 16 file .docx nghiệp vụ -> chunk theo heading/đoạn -> gắn nhãn category theo tên file
  - Đọc FAQ .xlsx (câu hỏi - trả lời, phân loại nhiều cấp) -> mỗi cặp Q&A là 1 knowledge unit,
    category = "faq", subcategory = nhãn phân loại trong file
  - Làm sạch: chuẩn hoá khoảng trắng, loại ký tự lỗi mã hoá phổ biến (mojibake), bỏ dòng rỗng
  - Dedup: exact-hash + near-dup do MemoryStore.add_knowledge() xử lý khi insert
"""
import re
import unicodedata
from pathlib import Path

from docx import Document
import openpyxl

import config
from memory_api import MemoryStore


# ---------------------------------------------------------------- cleaning
_MOJIBAKE_MAP = {
    "â€™": "’", "â€œ": "“", "â€": "”", "Ã¡": "á", "Ã©": "é",
    "Ã­": "í", "Ã³": "ó", "Ã º": "ú", "\ufeff": "",
}


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = unicodedata.normalize("NFC", text)
    for bad, good in _MOJIBAKE_MAP.items():
        text = text.replace(bad, good)
    text = text.replace("\x0c", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def guess_category(filename: str) -> str:
    name = filename.lower()
    rules = {
        "bao_hiem": ["bảo hiểm", "bao hiem", "insurance"],
        "bhxh": ["bhxh", "bảo hiểm xã hội"],
        "vay_tien": ["vay", "loan", "tín dụng", "tin dung"],
        "nap_chuyen_tien": ["nạp", "nap", "chuyển tiền", "chuyen tien", "liên ngân hàng", "lien ngan hang"],
        "cuoc_vien_thong": ["cước", "cuoc", "viễn thông", "vien thong", "topup", "sim"],
        "hoa_don_dien_nuoc": ["hóa đơn", "hoa don", "điện", "dien", "nước", "nuoc"],
        "tai_chinh": ["tài chính", "tai chinh", "finance"],
        "tong_quan": ["tổng quan", "tong quan", "giới thiệu", "gioi thieu", "overview"],
    }
    for cat, keywords in rules.items():
        if any(k in name for k in keywords):
            return cat
    return "khac"


# ---------------------------------------------------------------- chunking
def chunk_text(paragraphs: list[str], max_chars: int = None, overlap: int = None) -> list[str]:
    """Gộp các đoạn liên tiếp thành chunk <= max_chars, có overlap để không cắt mất ngữ cảnh."""
    max_chars = max_chars or config.CHUNK_MAX_CHARS
    overlap = overlap or config.CHUNK_OVERLAP_CHARS
    chunks, buf = [], ""
    for para in paragraphs:
        para = clean_text(para)
        if not para:
            continue
        if len(buf) + len(para) + 1 <= max_chars:
            buf = f"{buf}\n{para}".strip()
        else:
            if buf:
                chunks.append(buf)
            tail = buf[-overlap:] if overlap and buf else ""
            buf = f"{tail}\n{para}".strip()
    if buf:
        chunks.append(buf)
    return chunks


# ---------------------------------------------------------------- docx
def parse_docx(path: Path) -> list[dict]:
    """
    Trả về list[{content, source_ref}] theo cấu trúc heading của file .docx:
    mỗi heading (Heading 1/2/3) mở 1 section mới, nội dung dưới heading được chunk riêng,
    source_ref = đường dẫn heading (vd 'Chương 2 > Điều kiện vay').
    """
    doc = Document(str(path))
    sections = []
    current_heading_stack = []
    current_paras = []

    def flush():
        if current_paras:
            ref = " > ".join(current_heading_stack) if current_heading_stack else "(mở đầu)"
            for chunk in chunk_text(current_paras):
                sections.append({"content": chunk, "source_ref": ref})

    for para in doc.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            flush()
            current_paras = []
            level = 1
            m = re.search(r"(\d+)", style)
            if m:
                level = int(m.group(1))
            current_heading_stack = current_heading_stack[:level - 1] + [text]
        else:
            current_paras.append(text)
    flush()

    # bảng trong docx (nếu có) -> mỗi bảng thành 1 chunk dạng text có cấu trúc
    for ti, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            sections.append({"content": "\n".join(rows_text), "source_ref": f"Bảng {ti + 1}"})

    return sections


# ---------------------------------------------------------------- xlsx FAQ
def parse_faq_xlsx(path: Path, question_col: str = None, answer_col: str = None) -> list[dict]:
    """
    Đọc FAQ .xlsx. Tự dò cột câu hỏi/trả lời theo header (không phân biệt hoa/thường,
    chấp nhận tiếng Việt có/không dấu) nếu không truyền tên cột.
    Các cột còn lại (nếu có, vd 'Phân loại', 'Chuyên mục') được ghép thành subcategory.
    Trả về list[{content, subcategory, source_ref}]
    """
    wb = openpyxl.load_workbook(str(path), data_only=True)
    results = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(h).strip().lower() if h else "" for h in rows[0]]

        def find_col(keywords):
            for i, h in enumerate(header):
                if any(k in h for k in keywords):
                    return i
            return None

        q_idx = find_col(["câu hỏi", "cau hoi", "question"]) if question_col is None else header.index(question_col.lower())
        a_idx = find_col(["trả lời", "tra loi", "answer", "đáp án", "dap an"]) if answer_col is None else header.index(answer_col.lower())
        cat_idxs = [i for i, h in enumerate(header)
                    if i not in (q_idx, a_idx) and any(k in h for k in
                        ["phân loại", "phan loai", "chuyên mục", "chuyen muc", "category", "nhóm", "nhom"])]

        if q_idx is None or a_idx is None:
            continue  # sheet này không phải FAQ dạng Q/A -> bỏ qua

        for row_i, row in enumerate(rows[1:], start=2):
            if row_i > len(rows):
                break
            q = clean_text(str(row[q_idx])) if q_idx < len(row) and row[q_idx] else ""
            a = clean_text(str(row[a_idx])) if a_idx < len(row) and row[a_idx] else ""
            if not q or not a:
                continue
            subcats = [str(row[i]).strip() for i in cat_idxs if i < len(row) and row[i]]
            content = f"Câu hỏi: {q}\nTrả lời: {a}"
            results.append({
                "content": content,
                "subcategory": " / ".join(subcats) if subcats else None,
                "source_ref": f"{ws.title}!row{row_i}",
            })
    return results


# ---------------------------------------------------------------- orchestration
def ingest_docx_folder(folder: Path, store: MemoryStore, ttl_days: int = None) -> dict:
    folder = Path(folder)
    summary = {"files": 0, "inserted": 0, "exact_dup_skipped": 0, "near_dup_superseded": 0, "errors": []}
    for path in sorted(folder.glob("*.docx")):
        summary["files"] += 1
        category = guess_category(path.name)
        try:
            sections = parse_docx(path)
        except Exception as e:  # noqa: BLE001
            summary["errors"].append(f"{path.name}: {e}")
            continue
        for sec in sections:
            res = store.add_knowledge(
                content=sec["content"], category=category, source_file=path.name,
                source_ref=sec["source_ref"], ttl_days=ttl_days,
            )
            summary[res["status"]] = summary.get(res["status"], 0) + 1
    return summary


def ingest_faq_xlsx(path: Path, store: MemoryStore, ttl_days: int = None) -> dict:
    path = Path(path)
    summary = {"files": 1, "inserted": 0, "exact_dup_skipped": 0, "near_dup_superseded": 0, "errors": []}
    try:
        items = parse_faq_xlsx(path)
    except Exception as e:  # noqa: BLE001
        summary["errors"].append(f"{path.name}: {e}")
        return summary
    for item in items:
        res = store.add_knowledge(
            content=item["content"], category="faq", subcategory=item.get("subcategory"),
            source_file=path.name, source_ref=item["source_ref"], ttl_days=ttl_days,
        )
        summary[res["status"]] = summary.get(res["status"], 0) + 1
    return summary
